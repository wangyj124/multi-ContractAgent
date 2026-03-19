
import os
import shutil
from docx import Document

def generate_scenario_docs(output_dir: str = "data/input/scenarios"):
    """
    Generate specific DOCX files for different test scenarios.
    """
    if os.path.exists(output_dir):
        shutil.rmtree(output_dir)
    os.makedirs(output_dir, exist_ok=True)
    
    # Scenario 1: Normal Success
    # Doc has clear answer.
    doc1 = Document()
    doc1.add_paragraph("第一章 合同总价", style="Heading 1")
    doc1.add_paragraph("1.1 本合同总价为人民币100万元整。")
    doc1.save(os.path.join(output_dir, "scenario_1_success.docx"))
    
    # Scenario 2: Present but Empty
    # Doc mentions the field but value is blank.
    doc2 = Document()
    doc2.add_paragraph("第一章 关键信息", style="Heading 1")
    doc2.add_paragraph("1.1 签署日期：__________") # Blank
    doc2.add_paragraph("1.2 生效日期：(空)") # Explicitly empty
    doc2.save(os.path.join(output_dir, "scenario_2_empty.docx"))
    
    # Scenario 3: Not Found
    # Doc does not have the info at all.
    doc3 = Document()
    doc3.add_paragraph("第一章 双方义务", style="Heading 1")
    doc3.add_paragraph("1.1 甲乙双方应友好协商。")
    doc3.save(os.path.join(output_dir, "scenario_3_missing.docx"))
    
    # Scenario 4: Distributed / Ambiguous (Requires Retry/Context)
    # Answer might be split or first hit is partial.
    doc4 = Document()
    doc4.add_paragraph("第一章 付款方式", style="Heading 1")
    doc4.add_paragraph("1.1 预付款：合同签订后支付30%。")
    doc4.add_paragraph("（中间隔了很多无关内容...）")
    doc4.add_page_break()
    doc4.add_paragraph("第二章 补充条款", style="Heading 1")
    doc4.add_paragraph("2.1 关于付款的补充规定：剩余70%在验收合格后支付。")
    # Query: "分期付款详情" -> Worker might just find 1.1 first. Validator should ideally say "is that all?". 
    # But current Validator logic is mostly checking format/logic of extracted value.
    # To trigger retry, Validator needs to know there's more.
    # Or maybe we test "Conflicting info" where one place says X and another says Y?
    # Let's try "Partial info" -> validator prompt says "check for total sum=100%".
    # If worker only extracts 30%, validator says "Sum is 30%, expected 100%".
    doc4.save(os.path.join(output_dir, "scenario_4_retry.docx"))

    print(f"Generated 4 scenario documents in {output_dir}")

if __name__ == "__main__":
    generate_scenario_docs()
