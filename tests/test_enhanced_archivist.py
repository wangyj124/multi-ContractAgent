import unittest
from unittest.mock import MagicMock, patch
import tempfile
import os
from docx import Document
from src.core.archivist import Archivist

class TestEnhancedArchivist(unittest.TestCase):
    def setUp(self):
        self.archivist = Archivist()
        # Mock embedding model to avoid errors
        self.archivist.embedding_model = MagicMock(return_value=[0.1]*1024)
        # Mock _generate_summary to avoid real LLM calls if not patched in test
        self.archivist._generate_summary = MagicMock(return_value="Mock Summary")

    def test_detect_hierarchy_volume(self):
        """Test detection of Volume (Depth 1)"""
        current = []
        text = "第一卷"
        new_hierarchy = self.archivist._detect_hierarchy(text, current)
        self.assertEqual(new_hierarchy, ["第一卷"])

    def test_detect_hierarchy_chapter(self):
        """Test detection of Chapter (Depth 2)"""
        # Case 1: Start with empty -> Should pad with empty string for Volume
        current = []
        text = "第一章"
        new_hierarchy = self.archivist._detect_hierarchy(text, current)
        self.assertEqual(new_hierarchy, ["", "第一章"])

        # Case 2: Under Volume -> Should be [Volume, Chapter]
        current = ["第一卷"]
        text = "第二章"
        new_hierarchy = self.archivist._detect_hierarchy(text, current)
        self.assertEqual(new_hierarchy, ["第一卷", "第二章"])

    def test_detect_hierarchy_section(self):
        """Test detection of Section (Depth 3)"""
        # Section is Depth 3.
        # Current: [Volume(empty), Chapter]
        current = ["", "第一章"]
        text = "1.1 节"
        new_hierarchy = self.archivist._detect_hierarchy(text, current)
        self.assertEqual(new_hierarchy, ["", "第一章", "1.1 节"])

    def test_detect_hierarchy_dynamic(self):
        """Test dynamic depth calculation or standard patterns for X.X"""
        # 1.1 -> Depth 3 (Standard Section pattern)
        # Current: [Volume(empty), Chapter]
        current = ["", "第一章"]
        text = "1.1 Introduction"
        new_hierarchy = self.archivist._detect_hierarchy(text, current)
        self.assertEqual(new_hierarchy, ["", "第一章", "1.1 Introduction"])

        # 1.1.1 -> Depth 4 (Clause pattern)
        # Current: [Volume(empty), Chapter, Section]
        current = ["", "第一章", "1.1 Introduction"]
        text = "1.1.1 Details"
        new_hierarchy = self.archivist._detect_hierarchy(text, current)
        self.assertEqual(new_hierarchy, ["", "第一章", "1.1 Introduction", "1.1.1 Details"])

        # Back to 1.2 -> Depth 3
        # Should replace 1.1 Introduction
        current = ["", "第一章", "1.1 Introduction", "1.1.1 Details"]
        text = "1.2 Conclusion"
        new_hierarchy = self.archivist._detect_hierarchy(text, current)
        self.assertEqual(new_hierarchy, ["", "第一章", "1.2 Conclusion"])

    def test_detect_hierarchy_signature_page(self):
        """Test Signature Page (Depth 1)"""
        current = ["", "第一章", "1.1"]
        text = "此页为合同签字页"
        new_hierarchy = self.archivist._detect_hierarchy(text, current)
        self.assertEqual(new_hierarchy, ["此页为合同签字页"])

    @patch('src.core.archivist.get_llm')
    def test_full_extraction_flow_with_docx(self, mock_get_llm):
        """
        Test full extraction flow using a real temporary docx file.
        """
        # Setup Mock LLM
        mock_llm_instance = MagicMock()
        mock_llm_instance.invoke.return_value.content = "Mock Summary Content"
        mock_get_llm.return_value = mock_llm_instance

        # Create a fresh instance to ensure get_llm is mocked properly if needed
        # (Though we are patching the module function, so it should work)
        archivist = Archivist()
        archivist.embedding_model = MagicMock(return_value=[0.1]*1024)

        # Create a temporary docx file
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name
        
        try:
            doc = Document()
            # Chunk 1: Chapter 1 (Depth 2) -> Path: /Chapter 1
            doc.add_paragraph("第一章")
            doc.add_paragraph("Content of chapter 1.")
            
            # Chunk 2: Section 1 (under Chapter 1, Depth 3) -> Path: /Chapter 1/Section 1
            doc.add_paragraph("1.1 节")
            doc.add_paragraph("Content of section 1.")
            
            # Chunk 3: Table (under Section 1)
            table = doc.add_table(rows=1, cols=2)
            cell = table.cell(0, 0)
            cell.text = "Header 1"
            cell = table.cell(0, 1)
            cell.text = "Header 2"
            
            doc.save(tmp_path)
            
            # Run extraction
            chunks = archivist.extract_chunks(tmp_path)
            
            # Verify results
            self.assertTrue(len(chunks) >= 3, f"Expected at least 3 chunks, got {len(chunks)}")
            
            # 1. Verify Chunk 1 (Chapter 1)
            chunk1 = chunks[0]
            self.assertTrue(chunk1['content'].startswith("第一章"), f"Chunk 1 content: {chunk1['content']}")
            # Path should not contain empty strings
            # Hierarchy: ["", "第一章"] -> Path: "第一章"
            self.assertEqual(chunk1['metadata']['path'], "第一章")
            
            # 2. Verify Chunk 2 (Section 1)
            chunk2 = chunks[1]
            self.assertTrue(chunk2['content'].startswith("1.1 节"), f"Chunk 2 content: {chunk2['content']}")
            # Hierarchy: ["", "第一章", "1.1 节"] -> Path: "第一章/1.1 节"
            self.assertEqual(chunk2['metadata']['path'], "第一章/1.1 节")
            
            # 3. Verify Chunk 3 (Table)
            table_chunk = next((c for c in chunks if c['metadata']['type'] == 'table'), None)
            self.assertIsNotNone(table_chunk, "Should find a table chunk")
            self.assertEqual(table_chunk['metadata']['path'], "第一章/1.1 节")
            
            # Verify LLM was called
            self.assertTrue(mock_get_llm.called)
            
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)

    def test_generate_document_structure(self):
        """Test generating document structure from chunks with string paths"""
        chunks = [
            {'metadata': {'path': '第一章'}},
            {'metadata': {'path': '第一章/1.1 节'}},
            {'metadata': {'path': '第二章'}}
        ]
        
        structure = self.archivist.generate_document_structure(chunks)
        
        # Verify
        self.assertIn("第一章", structure)
        self.assertIn("  1.1 节", structure)
        self.assertIn("第二章", structure)
        
        # Verify no weird splits
        self.assertNotIn("  第", structure) 

if __name__ == '__main__':
    unittest.main()
