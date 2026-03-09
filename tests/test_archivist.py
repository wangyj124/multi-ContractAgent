import os
import sys
import pytest

# Add src to path if not already there
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '../')))

from src.core.archivist import Archivist

def test_archivist_initialization():
    archivist = Archivist()
    assert archivist.embedding_model_name == "qwen3-embedding-8B"

from docx import Document as DocxDocument

def create_dummy_docx(path):
    doc = DocxDocument()
    
    # Chapter 1
    doc.add_heading('Chapter 1: Introduction', level=1)
    doc.add_paragraph('This is the introduction text. It provides an overview of the document.')
    
    # 1.1
    doc.add_heading('1.1 Background', level=2)
    doc.add_paragraph('This section explains the background of the project.')
    
    # Long paragraph for chunking test
    long_text = " ".join(["Word"] * 1000)
    doc.add_paragraph(f"This is a long paragraph to test chunking. {long_text}")
    
    # 1.1.1
    doc.add_heading('1.1.1 Specific Details', level=3)
    doc.add_paragraph('Here are some specific details about the background.')
    
    # Table
    table = doc.add_table(rows=2, cols=2)
    cell = table.cell(0, 0)
    cell.text = 'Header 1'
    cell = table.cell(0, 1)
    cell.text = 'Header 2'
    cell = table.cell(1, 0)
    cell.text = 'Value 1'
    cell = table.cell(1, 1)
    cell.text = 'Value 2'
    
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)

def test_extract_chunks_from_dummy():
    dummy_path = "data/input/dummy.docx"
    if not os.path.exists(dummy_path):
        create_dummy_docx(dummy_path)
        
    archivist = Archivist()
    chunks = archivist.extract_chunks(dummy_path)
    
    assert isinstance(chunks, list)
    assert len(chunks) > 0
    
    # Check for expected hierarchy
    # We expect "Chapter 1: Introduction", "1.1 Background", "1.1.1 Specific Details"
    
    paths = [c["metadata"]["path"] for c in chunks]
    
    # Check if we have at least one chunk with "Chapter 1..."
    has_chapter_1 = any(len(p) >= 1 and "Chapter 1" in p[0] for p in paths)
    assert has_chapter_1, "Should detect Chapter 1"
    
    # Check if we have at least one chunk with "1.1..."
    has_section_1_1 = any(len(p) >= 2 and "1.1" in p[1] for p in paths)
    assert has_section_1_1, "Should detect 1.1"
    
    # Check if we have at least one chunk with "1.1.1..."
    has_section_1_1_1 = any(len(p) >= 3 and "1.1.1" in p[2] for p in paths)
    assert has_section_1_1_1, "Should detect 1.1.1"
    
    # Check chunking
    # Find the long paragraph chunk(s) under 1.1 (but before 1.1.1)
    # The long text is under 1.1 Background.
    long_section_chunks = [c for c in chunks if len(c["metadata"]["path"]) == 2 and "1.1" in c["metadata"]["path"][1]]
    
    # We expect multiple chunks if the text was long enough
    # In dummy docx: "Word" * 1000 ~ 5000 chars.
    # Archivist chunk size: 4000 chars.
    # So we expect at least 2 chunks.
    # Note: The logic in `extract_chunks` accumulates text until hierarchy changes.
    # So the whole 1.1 section (including the intro sentence + long text) is buffered.
    # Total length > 4000. So it should be chunked.
    
    # However, "1.1.1" starts a new hierarchy. So 1.1 buffer is flushed before 1.1.1.
    
    # Let's check count of chunks for 1.1 section
    assert len(long_section_chunks) >= 2, f"Long section should be chunked, got {len(long_section_chunks)} chunks"
    
    # Check table
    table_chunks = [c for c in chunks if c["metadata"]["type"] == "table"]
    assert len(table_chunks) > 0
    assert "Header 1" in table_chunks[0]["content"]

if __name__ == "__main__":
    pytest.main([__file__])
