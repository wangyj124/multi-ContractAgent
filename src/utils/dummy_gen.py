
import os
import random
from docx import Document
from typing import List

def generate_dummy_contract(output_path: str = "data/input/dummy_contract.docx"):
    """
    Generate a dummy contract DOCX file for testing.
    Includes Cover, TOC, Chapters, Clauses, and Tables.
    """
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    doc = Document()
    
    # Cover
    doc.add_paragraph("合同封面", style="Heading 1")
    doc.add_paragraph("合同编号：HT-2023-001")
    doc.add_paragraph("项目名称：测试项目")
    doc.add_page_break()
    
    # TOC (Simulated)
    doc.add_paragraph("目录", style="Heading 1")
    doc.add_paragraph("第一章 定义................................1")
    doc.add_paragraph("第二章 合同标的............................2")
    doc.add_paragraph("第三章 价格................................3")
    doc.add_page_break()
    
    # Volume 1
    doc.add_paragraph("第一卷 商务条款", style="Heading 1")
    
    # Chapter 1: Definitions
    doc.add_paragraph("第一章 定义", style="Heading 2")
    doc.add_paragraph("1.1 买方：指测试公司A。")
    doc.add_paragraph("1.2 卖方：指测试公司B。")
    doc.add_paragraph("1.3 合同设备：指附件1中列出的设备。")
    
    # Chapter 2: Scope
    doc.add_paragraph("第二章 合同标的", style="Heading 2")
    doc.add_paragraph("2.1 本合同标的为买方向卖方采购的设备。")
    doc.add_paragraph("2.2 卖方保证设备符合国家标准GB/T 12345-2008。")
    
    # Chapter 3: Price
    doc.add_paragraph("第三章 价格", style="Heading 2")
    doc.add_paragraph("3.1 合同总价")
    doc.add_paragraph("合同总价为人民币壹佰万元整（¥1,000,000.00）。")
    doc.add_paragraph("3.2 价格包含设备费、运输费、保险费及税费。")
    
    # Chapter 4: Payment
    doc.add_paragraph("第四章 支付条件", style="Heading 2")
    doc.add_paragraph("4.1 预付款")
    doc.add_paragraph("合同签署后10个工作日内，买方支付合同总价的30%作为预付款。")
    doc.add_paragraph("4.2 进度款")
    doc.add_paragraph("设备到货验收合格后，支付合同总价的60%。")
    doc.add_paragraph("4.3 质保金")
    doc.add_paragraph("质保期满且无质量问题后，支付剩余10%。")
    
    # Table example
    doc.add_paragraph("附件1：设备清单", style="Heading 2")
    table = doc.add_table(rows=4, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '设备名称'
    hdr_cells[1].text = '数量'
    hdr_cells[2].text = '单价'
    
    data = [
        ['服务器', '10', '50000'],
        ['交换机', '5', '20000'],
        ['存储柜', '2', '100000']
    ]
    
    for i, row_data in enumerate(data):
        row_cells = table.rows[i+1].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
            
    # Sign Page
    doc.add_page_break()
    doc.add_paragraph("此页为合同签字页", style="Heading 1")
    doc.add_paragraph("买方（盖章）：________________")
    doc.add_paragraph("卖方（盖章）：________________")
    doc.add_paragraph("签署日期：2023年10月1日")

    doc.save(output_path)
    print(f"Dummy contract generated at: {output_path}")

if __name__ == "__main__":
    generate_dummy_contract()
